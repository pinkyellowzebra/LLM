# -*- coding: utf-8 -*-
"""
Created on Fri July 26 16:13:23 2024
Create a Streamlit app where users sequentially input a YouTube video ID and a Gemini API key. The app will then provide a summary based on the video's title, description, and subtitles. After that, a session will be established in which users can engage in multiple rounds of questions and answers via chat.
@author: Hui, Hong
"""
# Module Import
import streamlit as st
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
import googleapiclient.discovery
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api.formatters import TextFormatter
import google.generativeai as genai
from iso639 import languages
import string
from langchain.callbacks.base import BaseCallbackHandler
from langchain.schema import ChatMessage
from io import BytesIO
from docx import Document

# API Key Validation
def valid_api_key(api_key):
    if len(api_key) > 64:
        return False
    
    valid_chars = string.ascii_letters + string.digits + "!@#$%^&*()_-+={}|[]:;'<>,.?/~"
    
    for char in api_key:
        if char not in valid_chars:
            return False
            
    return True

# Retrieve YouTube Video Info: Title, Description
def get_video_info(video_id, developer_key):
    title, description = "",""
    try:
        youtube = build(serviceName='youtube', version='v3', developerKey=developer_key)
        # Using videos().list to fetch video information
        video_response = youtube.videos().list(part='snippet',id=video_id).execute()
        # Fetch the title and description of the first video
        video = video_response['items'][0]
        title = video['snippet']['title']
        description = video['snippet']['description']
        
    except HttpError as e:
        st.error(f'Error fetching video info: {e}')
    
    return title, description 

# Retrieve ISO 639-1 Language Code
def get_iso_639_1_codes():
    iso_codes = []
    for lang in languages.part1:
        iso_codes.append(lang)
    iso_codes.append('en-US')
    iso_codes.append('en-GB')
    
    return iso_codes

# Retrieve YouTube Subtitles
def get_transcript(video_id, languages=get_iso_639_1_codes()):
    try:
        transcript=YouTubeTranscriptApi.get_transcript(video_id, languages=languages)
        transcript=TextFormatter().format_transcript(transcript)
    
    except Exception as e:
        st.error(f'Error fetching transcript:{e}')
    
    return ""

# Chat Session Management Class: Manages the chat interactions with the Gemini API
class ChatSession:
    # Configures the Gemini API with the provided API key and initializes the chat model
    def __init__(self, gemini_api_key, model_name="models/gemini-1.5-pro-latest"):
        # Configure the Gemini API key
        genai.configure(api_key=gemini_api_key)  
        # Initialize the model
        self.model = genai.GenerativeModel(model_name)
        self.history = []
        
    def send_message(self, question):
        # Create new messages, including previous conversation history
        #context = " ".join([f"{q} {a}" for q, a in self.history]) + " " + question
        context = question
        # Generate response from the model
        response = self.model.generate_content(context)

        # Extract text from the response object
        if response.candidates and response.candidates[0].content.parts:
            response_text = response.candidates[0].content.parts[0].text
        else:
            raise ValueError("No valid response received from the model.")
        # Append the question and response after the full response is received
        self.history.append((question, response_text))  
        
        return response_text
        
    def get_history(self):
        return self.history # Return all chat logs

    def clear_history(self):
        self.history = [] # Clear chat history

# Ftn to convert video summary(response1) results into a Word document
def download_to_word(text):
    doc = Document()
    doc.add_heading('Summary', 0)
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
    
# Stream Handler Class: Handles and displays streaming text content in the Streamlit app
class StreamHandler(BaseCallbackHandler):
    # Sets up a container for displaying streamed text and initializes an empty string for the text
    def __init__(self, container):
        self.container = container
        self.text = ""
    # Appends new tokens of text as they are received and updates the Streamlit UI in real-time  
    def new_token(self, token: str, **kwargs) -> None:
        self.text += token
        self.container.markdown(self.text)

# Main Ftn
def main():     
    # Start Streamlit app
    st.set_page_config(page_title="InfoFindr: Finding Information in YouTube Videos", page_icon="🔍")
    st.title("🔎 InfoFindr: Finding Information in YouTube Videos")
    
    video_id = st.text_input('Enter the YouTube video ID. the ID can be found after https://www.youtube.com/watch?v= in the link.')
    
    # Check user inputs
    if not video_id:
        st.info("Please enter the YouTube video ID to continue.")
    else:
        st.success("You have successfully entered the video ID.")
        gemini_api_key = st.sidebar.text_input(label="Gemini API Key", type="password")
        # Check user inputs
        if not gemini_api_key:
            st.info("Please add your Gemini API key to continue.") #gemini_api_key = 'AIzaSyAij_2eLy-nzQVRDouKKur-1TfObHYi3E8' 
        else:
            if not valid_api_key(gemini_api_key):
                st.sidebar.error("Invalid Gemini API key. Please check and try again.")
            else:
                st.sidebar.success("Gemini API key added successfully.")

                # Process entered video
                developer_key = 'AIzaSyA_1WmZ4E9fV3G1xFZyKJfwhLxWyaVnxUg'
                title, description = get_video_info(video_id, developer_key)
                transcript = get_transcript(video_id) 

                # Display video
                st.video(f"https://www.youtube.com/watch?v={video_id}")
                
                session = ChatSession(gemini_api_key)
                response1 = session.send_message(f"The title of the video is {title}, the description is {description}, and the transcript is {transcript}, but please summarize the entire script by correcting any misrepresented technical terms based on the given information.") 
                st.markdown(response1)

                # Download the results(response1)
                # Layout with columns to align the button to the right
                col1, col2 = st.columns([4, 1])  

                with col1:
                    st.write("")  # Placeholder for alignment

                with col2:
                    doc_buffer = download_to_word(response1)
                    st.download_button(
                        label="Download",
                        data=doc_buffer,
                        file_name="summary.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                if "messages" not in st.session_state:
                    st.session_state["messages"] = [ChatMessage(role="assistant", content="If you have any questions, please ask.")]

                # Display existing msgs
                for msg in st.session_state.messages:
                    st.chat_message(msg.role).write(msg.content)
                    
                # User input handling
                if prompt := st.chat_input():
                    st.session_state.messages.append(ChatMessage(role="user", content=prompt))
                    st.chat_message("user").write(prompt)
                    # Generate response 
                    response2 = session.send_message(prompt)
                    st.chat_message("assistant").write(response2)
                    st.session_state.messages.append(ChatMessage(role="assistant", content=response2))
                        
if __name__ == "__main__":
    main()
